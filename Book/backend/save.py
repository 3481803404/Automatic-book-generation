from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING


def save_docx(json_data):
    title = json_data.get('title')
    chapters = json_data.get('chapters')
    filepath = f"./docxs/{title}.docx"

    doc = Document()

    # 书籍标题
    doc.add_heading(title, level=0)

    # 书籍目录
    doc.add_page_break()
    doc.add_heading('目录', level=1)
    for idx, chapter in enumerate(chapters, start=1):
        doc.add_paragraph(f"{idx}. {chapter['title']}", style='TOCHeading')

    for chapter in chapters:
        doc.add_page_break()
        # 添加章节标题
        chapter_title = doc.add_heading(chapter['title'], level=1)
        chapter_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        chapter_title.runs[0].font.size = Pt(16)  # 设置字体大小为16pt
        chapter_title.paragraph_format.space_after = Pt(12)  # 设置段后间距为12pt

        # 添加章节内容，设置1.5倍行距
        content_paragraph = doc.add_paragraph(chapter['content'])
        content_paragraph_format = content_paragraph.paragraph_format
        content_paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 保存文件
    doc.save(filepath)
    return filepath
